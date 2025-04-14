# -*- coding: utf-8 -*-
import pypandoc
from docx.shared import Pt, RGBColor,Inches
from datetime import datetime
import requests
import os
import time
from app.io.session import engine
from sqlmodel import Session, select, update, func, or_
from app.model.formal_news import FormalNews
from io import BytesIO
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

# class data:
#     title = "Poll suspects local official involvement"
#     title_translate = "民意调查怀疑当地官员参与其中"
#     translate = '国家发展管理研究所的一项调查显示，大多数泰国人支持政府切断对缅甸的公用事业供应以打击呼叫中心帮派的决定，并认为这些非法行动得到了泰国官员的协助。Nida民意调查于2月10日至11日在全国1310名18岁及以上、具有不同教育水平和职业的人中进行。当被问及政府打击诈骗团伙的行动时，70.54%的人表示完全同意政府切断对缅甸诈骗分子活动地区的公用事业供应的举措。此外，21.07%的人 somewhat同意，5.34%的人 somewhat不同意，3.05%的人强烈反对。关于这些措施是否可能有效，60.92%的人表示在一定程度上有效，17.71%的人表示非常有效，15.95%的人表示效果甚微，5.42%的人认为它们根本不会有影响。关于泰国官员在缅甸勾结的传言，69.85%的人表示肯定有，26.87%的人不确定，只有3.28%的人表示确定没有泰国官员帮助呼叫中心诈骗分子。当被问及哪一组人更多——被诱骗为呼叫中心帮派工作的人还是自愿工作的人——答案如下：49.77%的人认为两组人数可能相等，25.80%的人表示大多数人自愿前往缅甸为诈骗分子工作。'
#     pic_set = "https://cdn01.allafrica.com/download/pic/main/main/csiid/00341892:f7e108e10ed96e488cd99f0a45babf92:arc614x376:w1470:us1.jpg"

# data_list = [data]*10

def add_hyperlink(paragraph, url, text, color, underline, size):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if not color is None:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Add underline if specified
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    
    # Add font size if specified
    if size is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))  # 字体大小单位为半磅，因此乘以 2
        rPr.append(sz)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

# 检索数据
def get_data_from_db(index_list):
    result = []
    with Session(engine, autoflush=False) as db:
        for i in index_list:
            smt = select(FormalNews).where(FormalNews.id == i)
            exist_data = db.exec(smt).one_or_none()
            if exist_data:
                result.append(exist_data)
    return result

# 内网上传版本
'''
    内网: 标题摘要不要链接 传图 10条左右 国际要闻
'''

# 按照政治军事类别对新闻分类，dict里套列表
def split_data(data_list):
    result = {}
    data_list = [i for i in data_list if i.main_classify in ['政治',"军事","社会","经济"]]
    # 处理
    for k in data_list:
        if k.main_classify in result.keys():
            result[k.main_classify].append(k)
        else:
            result[k.main_classify] = [k]
    new_result = {}
    if "政治" in result.keys():
        new_result["政治"] = result["政治"]
    
    if "军事" in result.keys():
        new_result["军事"] = result["军事"]
    
    if "社会" in result.keys():
        new_result["社会"] = result["社会"]
    
    if "经济" in result.keys():
        new_result["经济"] = result["经济"]
    return result

def split_data_by_main(data_list):
    result = {}
    final_result = {"国际要闻":[], "经济动态":[]}
    data_list = [i for i in data_list if i.main_classify in ['政治',"军事","社会","经济"]]
    # 处理
    for k in data_list:
        if k.main_classify in result.keys():
            result[k.main_classify].append(k)
        else:
            result[k.main_classify] = [k]
    
    # 遍历
    for k in ["政治","社会","军事"]:
        temp_data = result.get(k,[])
        final_result["国际要闻"].extend(temp_data)
    for k in ["经济"]:
        temp_data = result.get(k,[])
        final_result["经济动态"].extend(temp_data)
    # 去除空的
    final_result = {k:v for k,v in final_result.items() if len(v) > 0}
    return final_result

def inner_upload(origin_data, upload_pic_content):
    # 信息分类
    new_filter_data = split_data_by_main(origin_data)

    # 使用 pypandoc 转换 Markdown 为 Word 文档
    if not os.path.exists("temp_word"):
        os.mkdir("temp_word")
    file_name = "".join(str(datetime.now()).split(".")[0].replace("-", " ").replace(":"," ").split(" ")[:6])
    years = str(datetime.now()).split("-")[0]
    month = str(datetime.now()).split("-")[1]
    day = str(datetime.now()).split(" ")[0].split("-")[2]
    output_file = "每日外闻速览{}年{}月{}日-".format(years, month, day) + file_name +".docx"
    pypandoc.convert_text("", 'docx', format='md', outputfile=output_file)

    # 返回二进制word数据和文件名称
    byte_io = BytesIO()
    doc = Document(output_file)

    # 第一部分：标题
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("-每日外闻速览-")
    title_run.font.size = Pt(34)  # 字体大小设置为28磅
    title_run.font.name = '黑体'
    title_run.bold = True  # 加粗
    title_run.font.color.rgb = RGBColor(17, 76, 158)
    title_run.font.bold = True
    title_paragraph.alignment = 1  # 居中

    # 第二部分：增加日期描述
    years = str(datetime.now()).split(" ")[0].split("-")[0]
    month = str(datetime.now()).split(" ")[0].split("-")[1]
    day = str(datetime.now()).split(" ")[0].split("-")[2]

    normal_paragraph = doc.add_paragraph()
    normal_run = normal_paragraph.add_run("信息资源部{}年{}月{}日".format(years, month, day))
    normal_run.font.size = Pt(14)  # 字体大小设置为12磅
    normal_run.font.name = '黑体'
    normal_run.font.bold = True
    normal_paragraph.alignment = 1  # 居中

    # 增加分割线
    image_path = './element/pic1.png'  # 例如 'images/photo.jpg'
    pic_paragraph = doc.add_paragraph()
    pic_run = pic_paragraph.add_run()
    pic_run.add_picture(image_path, width=Inches(6))
    pic_paragraph.alignment = 1 

    # 第三部分：强调文本
    emphasis_paragraph = doc.add_paragraph()
    emphasis_run = emphasis_paragraph.add_run("声明：本文内容均直接采集自主要境外媒体，经过编译和整理。")
    emphasis_run.font.size = Pt(13)  # 字体大小设置为16磅
    emphasis_run.italic = True  # 设置为斜体
    emphasis_run.font.name = '等线'
    emphasis_run.font.color.rgb = RGBColor(169, 177, 184)
    emphasis_paragraph.paragraph_format.space_after = Pt(20)

    for topic, origin_data in new_filter_data.items():
        # 新增要闻展示
        image_path = './element/split.jpg'
        # 箭头+普通标题
        split_graph = doc.add_paragraph()
        split_graph.alignment = 1  # 居中
        split_graph.add_run().add_picture(image_path, width=Inches(0.1))
        split_run = split_graph.add_run(topic)
        split_run.font.size = Pt(18)
        split_run.font.bold = True
        split_run.font.name = '黑体'
        split_graph.add_run().add_picture(image_path, width=Inches(0.1))

        # 第四部分：添加
        for i in range(len(origin_data)):
            temp_data = origin_data[i]
            index = str(i+1)

            image_path = './element/tag.jpg'
            # 箭头+普通标题
            title_graph = doc.add_paragraph()
            title_graph.add_run().add_picture(image_path, width=Inches(0.25))
            title_run = title_graph.add_run("{}.{}".format(index, temp_data.title_translate))
            title_run.font.size = Pt(15)
            
            title_run.font.bold = True
            title_run.font.name = '黑体'
            title_run.line_spacing = Pt(21)

            # # 内容
            # content_paragraph = doc.add_paragraph()
            # # content_paragraph.paragraph_format.space_before = Pt(15)

            temp_ccc = str(temp_data.abstract).replace("\n","\n").split("\n")
            for temp_i in range(len(temp_ccc)):
                # 段前
                temp_graph = doc.add_paragraph()
                temp_c_run = temp_graph.add_run(temp_ccc[temp_i])
                temp_graph.paragraph_format.space_before = Pt(15)
                temp_c_run.font.size = Pt(15)  # 设置字体大小为16磅
                temp_c_run.font.name = '黑体'  # 设置字体为黑体.
                
                temp_graph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 设置为固定行距
                temp_graph.paragraph_format.line_spacing = Pt(30)  # 设置行距为 30 磅

                temp_graph.paragraph_format.space_after = Pt(15)

            # content_run = content_paragraph.add_run(str(temp_data.abstract).replace("\n","\n"))
            # content_run.font.size = Pt(15)
            # content_paragraph.paragraph_format.space_before = Pt(15)  # 段前间距
            # content_paragraph.paragraph_format.space_after = Pt(20)   # 段后间距

            # content_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 设置为固定行距
            # content_paragraph.paragraph_format.line_spacing = Pt(30)  # 设置行距为 30 磅
            
            # 图片下载
            state = 0
            pic_path = ""
            if temp_data.pic_set:
                url = temp_data.pic_set
                if ".gif" in url:
                    state = 0
                    pic_path = ""
                else:
                    try:
                        if not os.path.exists("temp_pic"):
                            os.mkdir("temp_pic")
                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/134.0.0.0 Safari/537.36'
                        }

                        temp_content = requests.get(url, timeout=5, headers=headers).content
                        c_pic_path = "temp_pic/"+str(time.time()).split(".")[0]+".jpg"
                        with open(c_pic_path,"wb") as f:
                            f.write(temp_content)
                        state = 1
                        pic_path = c_pic_path
                    except:
                        state = 0
                        pic_path = ""
            if state:
                try:
                    pic_paragraph = doc.add_paragraph()
                    pic_run = pic_paragraph.add_run()
                    pic_run.add_picture(pic_path, width=Inches(6))
                    os.remove(pic_path)
                except:
                    continue
            split_paragraph = doc.add_paragraph()
            split_paragraph.add_run("\n")
    # 市场数据
    image_path = './element/split.jpg'
    f_split_graph = doc.add_paragraph()
    f_split_graph.alignment = 1  # 居中
    f_split_graph.add_run().add_picture(image_path, width=Inches(0.1))
    f_split_run = f_split_graph.add_run("市场数据")
    f_split_run.font.size = Pt(16)
    f_split_run.font.bold = True
    f_split_run.font.name = '黑体'
    f_split_graph.add_run().add_picture(image_path, width=Inches(0.1))

    # 处理上传图片
    upload_pic_path = "temp_pic/"+str(time.time()).split(".")[0]+".jpg"
    with open(upload_pic_path, "wb") as f:
        f.write(upload_pic_content.getvalue())
    
    # 添加到最后
    if os.path.exists(upload_pic_path):
        pic_paragraph = doc.add_paragraph()
        pic_run = pic_paragraph.add_run()
        pic_run.add_picture(upload_pic_path, width=Inches(6))

    # 保存修改后的 Word 二进制数据
    doc.save(byte_io)
    byte_io.seek(0)
    if os.path.exists(output_file):
        os.remove(output_file)
    return byte_io, output_file

# 外网上传版本
'''
    外网: 标题和链接 20条左右 政治要闻、军事要闻
'''
def outter_upload(origin_data, upload_pic_content, ctype="国外域名"):
    if str(ctype) == "国外域名":
        link_base = "http://news.ideachorus.com/index.html?id={}"
    else:
        link_base = "http://150.158.25.36:8888/index.html?id={}"
    # 信息分类
    new_filter_data = split_data(origin_data)

    # 使用 pypandoc 转换 Markdown 为 Word 文档
    if not os.path.exists("temp_word"):
        os.mkdir("temp_word")
    file_name = "".join(str(datetime.now()).split(".")[0].replace("-", " ").replace(":"," ").split(" ")[:6])
    output_file = "外网新闻-"+ file_name +".docx"
    pypandoc.convert_text("", 'docx', format='md', outputfile=output_file)

    # 返回二进制word数据和文件名称
    byte_io = BytesIO()
    doc = Document(output_file)

    # 第一部分：标题
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("@每日外闻速览")

    # 第二部分：增加日期描述
    years = str(datetime.now()).split(" ")[0].split("-")[0]
    month = str(datetime.now()).split(" ")[0].split("-")[1]
    day = str(datetime.now()).split(" ")[0].split("-")[2]

    normal_paragraph = doc.add_paragraph()
    normal_run = normal_paragraph.add_run("*信息资源部{}年{}月{}日".format(years, month, day))

    # 第三部分：强调文本
    emphasis_paragraph = doc.add_paragraph()
    emphasis_run = emphasis_paragraph.add_run(">声明：本文内容均直接采集自主要境外媒体，点击标题查看原文。")
    
    for topic, data_list in new_filter_data.items():
        split_graph = doc.add_paragraph()
        split_run = split_graph.add_run("#"+str(topic) + "要闻")

        tag_graph = doc.add_paragraph()
        tag_run = tag_graph.add_run("##")

        # 第四部分：添加
        for i in range(len(data_list)):
            temp_data = data_list[i]
            index = str(i+1)
            t_graph = doc.add_paragraph()
            t_run = t_graph.add_run("{}.{}".format(index, temp_data.title_translate))

            p_graph = doc.add_paragraph()
            link = link_base.format(temp_data.id)
            add_hyperlink(p_graph, link, link,'000000', True, 12)
            p_run = p_graph.add_run()

            # 图片下载
            state = 0
            pic_path = ""
            if temp_data.pic_set:
                url = temp_data.pic_set
                if ".gif" in url:
                    state = 0
                    pic_path = ""
                else:
                    try:
                        if not os.path.exists("temp_pic"):
                            os.mkdir("temp_pic")
                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/134.0.0.0 Safari/537.36'
                        }

                        temp_content = requests.get(url, timeout=5, headers=headers).content
                        c_pic_path = "temp_pic/"+str(time.time()).split(".")[0]+".jpg"
                        with open(c_pic_path,"wb") as f:
                            f.write(temp_content)
                        state = 1
                        pic_path = c_pic_path
                    except Exception as e:
                        state = 0
                        pic_path = ""
            if state:
                try:
                    pic_paragraph = doc.add_paragraph()
                    pic_run = pic_paragraph.add_run()
                    pic_run.add_picture(pic_path, width=Inches(6))
                    os.remove(pic_path)
                except:
                    continue
    # 市场数据
    f_split_graph = doc.add_paragraph()
    f_split_run = f_split_graph.add_run("#市场数据")

    # 处理上传图片
    upload_pic_path = "temp_pic/"+str(time.time()).split(".")[0]+".jpg"
    with open(upload_pic_path, "wb") as f:
        f.write(upload_pic_content.getvalue())
    
    # 添加到最后
    if os.path.exists(upload_pic_path):
        pic_paragraph = doc.add_paragraph()
        pic_run = pic_paragraph.add_run()
        pic_run.add_picture(upload_pic_path, width=Inches(6))

    # 保存修改后的 Word 二进制数据
    doc.save(byte_io)
    byte_io.seek(0)

    if os.path.exists(output_file):
        os.remove(output_file)
    return byte_io, output_file